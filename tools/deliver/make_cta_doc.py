"""make_cta_doc.py, the .docx behind a reel's "comment X and I'll send you Y".

THE DELIVERABLE RULE. If the script says comment a keyword and you will send
something, that something is a REQUIRED deliverable at FIRST delivery, alongside
the MP4 and the caption. Not later, not when the owner asks. The reel starts collecting
comments the hour it posts and there is nothing to send.

WHAT THE DOC IS FOR, and this is the part that takes the time: it is where the
reel's overstatements get corrected. A hook is allowed to be a hook. The
document is what the viewer actually acts on, so it carries the caveats the
40 seconds could not:

  vid64  the blog said `masonry init claude-code`; the CLI itself marks init
         deprecated. The doc taught the current command.
  vid65  the MCP server needs Bun. The reel never says so, and it is the single
         most likely reason someone runs the six commands and gets nothing.
  vid66  the skill is a repo you clone, not a marketplace install.
  vid67  the VO says "without ever touching the terminal" and the repo's own
         Quickstart is three terminal commands. The doc says so in the first
         section rather than burying it, and adds that runs cost money.

Every claim in a doc is read out of the primary source (the repo via the GitHub
API, the CLI's own --help, the product's docs) on the day it is written, never
from memory, never from the reference video, never from third-party posts. Put
that provenance in the content JSON's `sources` field. It prints nowhere: it is
there so the next person knows what was verified and when.

House format, locked across the portfolio: US Letter, 1.25in side margins,
Calibri 11pt body, 28pt bold title, 16pt bold terracotta headings, Consolas
10.5pt for code, terracotta #B24A32 (the palette terracotta darkened for print).

Usage
-----
    python3 tools/deliver/make_cta_doc.py content.json -o out/vid67-agent.docx

On this machine python-docx needs the homebrew expat shim:
    DYLD_LIBRARY_PATH=/opt/homebrew/opt/expat/lib python3 tools/deliver/make_cta_doc.py ...

See tools/deliver/README.md and cta.example.json for the content schema.
"""
import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("python-docx is not installed. pip install python-docx")


INK = RGBColor(0x1A, 0x1A, 0x17)
MUTE = RGBColor(0x6E, 0x68, 0x5C)
ACC = RGBColor(0xB2, 0x4A, 0x32)      # terracotta, darkened for print
COLORS = {"ink": INK, "mute": MUTE, "accent": ACC}


class Doc:
    def __init__(self):
        self.d = Document()
        s = self.d.sections[0]
        s.page_width, s.page_height = Inches(8.5), Inches(11)
        s.left_margin = s.right_margin = Inches(1.25)
        s.top_margin = s.bottom_margin = Inches(1)
        n = self.d.styles["Normal"]
        n.font.name = "Calibri"
        n.font.size = Pt(11)
        n.font.color.rgb = INK
        n.paragraph_format.space_after = Pt(8)
        n.paragraph_format.line_spacing = 1.15

    def para(self, text="", size=11, bold=False, italic=False, color=INK,
             after=8, font="Calibri", style=None, indent=None, align=None):
        p = self.d.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(after)
        if indent is not None:
            p.paragraph_format.left_indent = Inches(indent)
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name, r.font.size = font, Pt(size)
        r.font.bold, r.font.italic, r.font.color.rgb = bold, italic, color
        return p

    def rich(self, parts, size=11, after=8, style=None):
        """A paragraph with a bold lead-in. The caveats read as caveats only
        when the first three words are bold: "You will use the terminal."
        Parts are ["text", bold] or ["text", bold, "accent"]."""
        p = self.d.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(after)
        for part in parts:
            text = part[0]
            bold = bool(part[1]) if len(part) > 1 else False
            col = COLORS.get(part[2], INK) if len(part) > 2 else INK
            r = p.add_run(text)
            r.font.name, r.font.size = "Calibri", Pt(size)
            r.font.bold, r.font.color.rgb = bold, col
        return p

    def code(self, lines):
        for ln in lines:
            p = self.d.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run(ln)
            r.font.name, r.font.size = "Consolas", Pt(10.5)
            r.font.color.rgb = INK

    def table(self, header, rows, widths=None, mono_col=None, after=10):
        t = self.d.add_table(rows=0, cols=len(header))
        t.style = "Table Grid"
        cells = t.add_row().cells
        for i, htxt in enumerate(header):
            r = cells[i].paragraphs[0].add_run(htxt)
            r.font.name, r.font.size, r.font.bold = "Calibri", Pt(10), True
            r.font.color.rgb = ACC
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                mono = mono_col is not None and i == mono_col
                r = cells[i].paragraphs[0].add_run(val)
                r.font.name = "Consolas" if mono else "Calibri"
                r.font.size = Pt(9.5 if mono else 10)
                r.font.color.rgb = INK
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        self.para("", after=after)
        return t


def build(content, out):
    d = Doc()
    d.para(content["title"], size=28, bold=True, after=2)
    if content.get("subtitle"):
        d.para(content["subtitle"], size=11, color=MUTE, after=16)

    for blk in content["blocks"]:
        k = blk["type"]
        if k == "h":
            d.para(blk["text"], size=16, bold=True, color=ACC, after=6)
        elif k == "h3":
            d.para(blk["text"], size=12, bold=True, after=4)
        elif k == "para":
            d.para(blk["text"], after=blk.get("after", 8),
                   italic=blk.get("italic", False),
                   color=COLORS.get(blk.get("color", "ink"), INK))
        elif k == "rich":
            d.rich(blk["parts"], after=blk.get("after", 8))
        elif k == "bullet":
            if blk.get("parts"):
                d.rich(blk["parts"], after=blk.get("after", 5), style="List Bullet")
            else:
                d.para(blk["text"], after=blk.get("after", 5), style="List Bullet")
        elif k == "step":
            d.para(blk["text"], after=blk.get("after", 4), style="List Number")
        elif k == "code":
            d.code(blk["lines"])
            d.para("", after=blk.get("after", 6))
        elif k == "table":
            d.table(blk["header"], blk["rows"], blk.get("widths"),
                    blk.get("mono_col"), blk.get("after", 10))
        elif k == "space":
            d.para("", after=blk.get("after", 8))
        else:
            sys.exit("unknown block type: %s" % k)

    if content.get("footer_link"):
        d.para(content["footer_link"], size=11, bold=True, color=ACC, after=2)
    if content.get("footer_note"):
        # licence, date created, date verified. The verification date is the
        # point: a doc that was true in June is a support ticket in August.
        d.para(content["footer_note"], size=9.5, color=MUTE)

    os.makedirs(os.path.dirname(os.path.abspath(out)), exist_ok=True)
    d.d.save(out)
    return out


def main():
    ap = argparse.ArgumentParser(description="build the CTA .docx from a content JSON")
    ap.add_argument("content", help="the film's CTA content JSON")
    ap.add_argument("-o", "--out", help="output .docx, default from the JSON's `out`")
    a = ap.parse_args()
    with open(a.content) as fh:
        content = json.load(fh)
    out = a.out or content.get("out")
    if not out:
        sys.exit("no output path. Pass -o, or set `out` in the JSON.")
    if not content.get("keyword"):
        # The keyword is what the viewer comments. A doc that does not name it
        # cannot be matched to the reel that promises it.
        print("WARNING: no `keyword` in the content JSON. Which comment does this answer?")
    print("wrote", build(content, out))


if __name__ == "__main__":
    main()
